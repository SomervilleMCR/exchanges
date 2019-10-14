"""Script for choosing guests on MCR exchanges."""
import numpy as np
import random
from csv import reader
import sys
import platform
from docx import Document
from docx.shared import Inches


#
# FUNCTION DEFINITIONS
#
def both_legs(x):
    """When someone can attend both legs of the exchange."""
    if (x[0] + x[1] == 2):
        return 1
    else:
        return 0


def any_legs(x):
    """When someone can attend one leg of the exchange."""
    if (x[0] + x[1] > 0):
        return 1
    else:
        return 0


#
# SCRIPT
#
if(len(sys.argv) != 2):
    filename = 'exchange_dinner_test_responses.csv'
    print('No valid exchange diners csv provided.')
    print('Will continue with the example file which is:')
    print(filename)
elif len(sys.argv) == 2:
    filename = sys.argv[1]
    print(filename)
exchange_titles = {'Hertford', 'Jesus'}

diners_dict = {}
for dinner in exchange_titles:
    print('Diners for ' + dinner)
    names = []
    home = []
    away = []
    emails = []
    dietary = []
    never = []
    with open(filename, 'r') as f:
        # SKIP OVER THE FIRST ROW OF THE SPREADSHEET
        next(f)
        # FOR EACH REMAINING ROW IN THE SPREADSHEET, EXTRACT THE INFORMATION
        for row in reader(f):
            names.append(row[1])
            emails.append(row[2])
            never.append("No" in row[-2])
            dietary.append(row[-3])
            attendance = row[3]
            home.append(dinner + ' Home' in attendance)
            away.append(dinner + ' Away' in attendance)

    # THIS SCRIPT WAS ORIGINALLY WRITTEN IN PYTHON 2. HOWEVER, IN PYTHON 3, THE
    # BEHAVIOUR OF THE zip() AND map() FUNCTIONS CHANGES. THE CODE BELOW TAKES
    # THIS INTO ACCOUNT AND IMPLEMENTS THE OLD PYTHON 2 BEHAVIOUR EVEN IF IT IS
    # RUN IN PYTHON 3.
    python_version = int(platform.python_version()[0])
    if python_version == 2:
        legs = zip(home, away)
        both = np.array(map(both_legs, legs))
        skip = np.array(map(any_legs, legs))
    elif python_version == 3:
        legs = []
        both = []
        skip = []
        for i, v in enumerate(home):
            legs.append((home[i], away[i]))
            both.append(both_legs([home[i], away[i]]))
            skip.append(any_legs([home[i], away[i]]))
        both = np.array(both)
        skip = np.array(skip)
    else:
        raise ValueError('The Python version is neither 2 nor 3')
    both_score = 1
    never_score = 5
    scores = np.array(skip) * (1 + both_score * np.array(both) +
                               never_score * np.array(never))
    scores = scores.tolist()
    for i in range(len(names)):
        print((names[i], scores[i], emails[i],
               never[i], home[i], away[i], dietary[i]))
    print('')

    # PUT ALL THE NAMES INTO A VIRTUAL BAG
    bag = []
    bounds = []
    total = 0
    for n in range(len(names)):
        x = scores[n]
        bag += [names[n]] * x
        bounds = bounds + [total + x]
        total += x

    # DRAW THE NAMES OF THE DINERS OUT OF THE VIRTUAL BAG
    diners = []
    while len(bag) > 0:
        draw = random.randint(1, len(bag)) - 1
        diners = diners + [bag[draw]]
        bag[:] = [x for x in bag if x != diners[-1]]

    # PRINT THE NAMES OF THE DINERS YOU HAVE DRAWN
    num_scores = len([s for s in scores if s != 0])
    print('Order for ' + dinner + ' (' + str(num_scores) + ')')
    print('name,never,home,away,dietary')
    for d in diners:
        index = names.index(d)
        print(d +
              ',' + str(scores[index]) +
              ',' + str(never[index]) +
              ',' + ('Home' if home[index] else '') +
              ',' + ('Away' if away[index] else '') +
              ',' + dietary[index])
    print('\n\n')

    # EXPORT THE NAMES OF THE PEOPLE WHO HAVE MADE IT ONTO THE EXCHANGE READY
    # TO BE COPY-PASTED INTO AN EMAIL
    with open('{}.txt'.format(dinner), 'w') as f:
        for diner in diners:
            f.write(diner + '\n')

    # SAVE NAMES OF DINERS TO A DICTIONARY
    diners_dict[dinner] = diners
    print(diners_dict[dinner])

# EXPORT TEXT TO A WORD DOCUMENT
document = Document()
document.add_heading('Trinity Exchange Dinner Lottery Results', level=1)
text = """
Hi Everyone,

Below is the list of diners for the two sets of exchange dinners this term. \
We've listed the members who won each lottery, so if you signed up for a \
dinner and don't see your name, then you are on the waiting list and might be \
contacted in the event of a dropout -- so stay tuned! Everyone on either of \
the lists below, please recheck your diaries and let me know if you can no \
longer make the dinners for any reason.

If your name is among the fifteen diners for either dinner, you are now \
responsible for paying £24 if you are attending both legs, or £12 if you are \
attending only one leg. You have until """
p = document.add_paragraph(text)
p.add_run('23:59').bold = True
p.add_run(' on ')
p.add_run('Wednesday, 1st May').bold = True
text = """, to pay for the dinners. It is preferred that you pay online to the \
MCR account (40-35-34, 33936759), but if you are unable to do that, you can \
pay person with cash to any executive member, or with cash in a sealed \
envelope with your name to the pidge of """
p.add_run(text)
p.add_run('Pippa (Philippa) Gleave').bold = True
p.add_run(', ')
p.add_run('Chris Whiteman').bold = True
p.add_run(' or ')
p.add_run('Laurel Kaye').bold = True
text = """ (let us know when you deliver it!). If you fail to make \
arrangements to pay for your place by the deadline, it will be forfeited to \
the next person on the waiting list.

If you previously signed up to a college formal on the same day as one of the \
exchanges you are going on """
p.add_run(text)
p.add_run('do not forget').bold = True
p.add_run(' to ')
p.add_run('cancel').underline = True
text = """ that spot by emailing catering at catering@some.ox.ac.uk to save \
you from paying for two tickets.

If, after paying for your spot on an exchange, you can no longer attend, \
simply contact us and we'll put you in touch with the next person on the \
waiting list. We can help with communications and arrangements, but it is \
ultimately your responsibility to arrange the swap and payment.

Once you've paid you can look forward to more information regarding each \
dinner as they approach.

Looking forward to the exchanges!

~ TripleSec

"""
p.add_run(text)
for exchange in exchange_titles:
    p.add_run(f'{exchange}\n').bold = True
    for diner in diners_dict[exchange]:
        p.add_run(f'{diner}\n')
    p.add_run('\n')
document.save('Trinity.docx')
